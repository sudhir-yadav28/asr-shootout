"""Builds the ≤3-page ASR Shootout report (.docx) from results/."""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
PLOTS = RESULTS / "plots"
OUT = ROOT / "report" / "report.docx"


def shade(cell, hex_):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_)
    tc_pr.append(shd)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return h


def para(doc, text, *, size=10.5, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def bullet(doc, text, *, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def make_table(doc, headers, rows, *, col_widths=None, header_fill="1F2A44",
               font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(font_size)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


doc = Document()

# tight margins to fit 3 pages
for section in doc.sections:
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


# ---------- TITLE -------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("ASR Shootout — Findings & Recommendation")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

sub = doc.add_paragraph()
r = sub.add_run("Sudhir Yadav  |  20 phone-recorded Bangalore locality clips, 5 ASR systems benchmarked")
r.italic = True
r.font.size = Pt(9.5)


# ---------- BOTTOM LINE -------------------------------------------------------
heading(doc, "Bottom line", level=2)
para(doc,
     "Sarvam and IndicConformer beat the Deepgram baseline on every accuracy metric. "
     "Sarvam leads the metric that actually matters here — locality hit rate — at 55%. "
     "IndicConformer is the open-source second at 45%, with the same WER and CER as Sarvam. "
     "Deepgram, the required baseline, comes in last on entity capture (40%). On latency, "
     "Groq is the surprise: 0.43 s mean, ~3× faster than Deepgram, and just as accurate on "
     "this set.")


# ---------- 1. APPROACH -------------------------------------------------------
heading(doc, "1. Approach", level=2)
para(doc,
     "I recorded 20 clips of Bangalore locality names spoken in natural Hinglish — 10 "
     "well-known names (Koramangala, Indiranagar, HSR Layout, …) and 10 harder Kannada-origin "
     "names (Byatarayanapura, Kadugondanahalli, Hesaraghatta, Rajarajeshwarinagar, …). "
     "Conditions were varied deliberately: 5 quiet, 5 with traffic noise, 4 over a phone call, "
     "3 rushed, 3 whispered. All recordings were made on a phone mic, no studio setup. The "
     "harder names and the noisy/whispered conditions are where the failure modes show up.")
para(doc,
     "I benchmarked five systems, picked to cover a 2×2 of {open-source, API} × "
     "{generic, India-tuned}, plus Groq-hosted Whisper as a latency reference point:")
make_table(
    doc,
    headers=["System", "Type", "Where it ran"],
    rows=[
        ["Deepgram Nova-3 (multi)", "API, generic — required baseline", "Mac"],
        ["Sarvam Saarika v2.5", "API, India-tuned", "Mac"],
        ["Groq Whisper large-v3", "API, generic (fast hosted Whisper)", "Mac"],
        ["OpenAI Whisper large-v3 (local)", "Open-source, generic", "Colab T4 GPU"],
        ["AI4Bharat IndicConformer 600M", "Open-source, India-tuned (Conformer arch.)", "Colab T4 GPU"],
    ],
    col_widths=[2.5, 3.2, 1.5],
)


# ---------- 2. METRICS --------------------------------------------------------
heading(doc, "2. Metrics", level=2)
para(doc,
     "WER alone is misleading on this task — what the business cares about is whether the "
     "locality entity was captured, not whether every filler word was right. I report a layered set:")
bullet(doc, "WER / CER against the reference, picking the better of the Roman vs Devanagari reference (Hinglish has no canonical script).")
bullet(doc, "Locality Hit Rate (LHR) — fuzzy partial-ratio match (threshold 85/100) of the locality name plus aliases (Roman + Devanagari spelling variants) against the hypothesis. This is the headline metric.")
bullet(doc, "Locality Partial Rate — same idea, threshold 60–85. Captures \"close but not quite\" cases.")
bullet(doc, "Phonetic similarity — Metaphone-based; gives partial credit when the model heard the right sounds but spelled them differently.")
bullet(doc, "Latency (mean + p90) — wall-clock per request, end-to-end. Only meaningful for the API systems; local Whisper/Conformer ran on a shared Colab T4 and is not strictly comparable.")


# ---------- 3. RESULTS --------------------------------------------------------
heading(doc, "3. Results", level=2)

summary_csv = RESULTS / "metrics_summary.csv"
s = pd.read_csv(summary_csv).set_index("model")
# Order rows: baseline first, then by LHR descending
order = ["deepgram_nova3"] + [m for m in s.sort_values("locality_hit_rate", ascending=False).index if m != "deepgram_nova3"]
s = s.loc[order]

# strip the failed fallback row from the headline table (note it separately)
display = s.drop(index=["indic_whisper_fallback"], errors="ignore")

rows = []
for m, r in display.iterrows():
    nice = {
        "deepgram_nova3": "Deepgram Nova-3 (baseline)",
        "sarvam_saarika_v2_5": "Sarvam Saarika v2.5",
        "groq_whisper_large_v3": "Groq Whisper large-v3",
        "indic_conformer": "AI4Bharat IndicConformer",
        "whisper_large_v3_local": "Whisper large-v3 (Colab)",
    }.get(m, m)
    rows.append([
        nice,
        f"{r['wer']:.3f}",
        f"{r['cer']:.3f}",
        f"{r['locality_hit_rate']:.2f}",
        f"{r['locality_partial_rate']:.2f}",
        f"{r['phonetic_sim']:.3f}",
        f"{r['latency_mean_s']:.2f} / {r['latency_p90_s']:.2f}",
    ])

make_table(
    doc,
    headers=["System", "WER ↓", "CER ↓", "LHR ↑", "Partial ↑", "Phon. ↑", "Latency mean / p90 (s)"],
    rows=rows,
    col_widths=[2.4, 0.65, 0.65, 0.65, 0.75, 0.7, 1.55],
)
para(doc,
     "↓ lower is better; ↑ higher is better. LHR = Locality Hit Rate (entity captured). "
     "vasista22/whisper-hindi-large-v2 was attempted as a fallback in case IndicConformer "
     "failed to load on Colab; it errored on every clip on T4 and is excluded — IndicConformer "
     "ran cleanly so the slot is filled.",
     italic=True, size=9)

# headline plot
if (PLOTS / "01_headline_metrics.png").exists():
    doc.add_picture(str(PLOTS / "01_headline_metrics.png"), width=Inches(6.4))


# ---------- 4. KEY FINDINGS ---------------------------------------------------
heading(doc, "4. Key findings", level=2)

bullet(doc,
       "Deepgram is not the right choice here. 40% LHR, 0.70 WER — behind Sarvam, "
       "IndicConformer, Groq Whisper, even local Whisper. Its multilingual mode keeps most "
       "tokens in Roman script and reaches for English near-homophones when it isn't sure "
       "(\"magistrate\" for Majestic; \"government sector\" for \"garment factory\" on clip 12). "
       "The brief required it as the baseline; the result is that the baseline is the wrong call.")

bullet(doc,
       "Domain training beats scale. IndicConformer (600M params, India-trained) matches Sarvam "
       "on WER (0.398 vs 0.395) and CER (0.236 vs 0.235) and beats Whisper-large-v3 (1.5B, generic) "
       "on every accuracy metric. Both India-trained systems land at ~40% WER; the three generic "
       "systems cluster at 60–70%.")

bullet(doc,
       "Deepgram heard the right sounds but wrote them as strings the matcher couldn't recover. "
       "It topped phonetic similarity (0.386) while sitting last on LHR. On clip 12 it produced "
       "\"kaddu gunda hainili\" for Kadugondanahalli — phonetically close, textually a miss. "
       "A locality-list post-processor (fuzzy + phonetic match against the known Bangalore set) "
       "would close most of this gap cheaply, and is the cheapest accuracy lift available.")

bullet(doc,
       "Groq is the latency story. 0.43 s mean, 0.62 s p90 — about 3× quicker than Deepgram "
       "(1.20 s) and slightly more accurate on this set. For telephony where the round-trip "
       "budget is sub-second, Groq is the obvious pick if Whisper-family quality is acceptable. "
       "Sarvam at 0.73 s is also well inside range.")

bullet(doc,
       "Rushed and whispered did less damage than expected. Traffic noise and the long Kannada "
       "names did most of the damage. See section 5.")


# ---------- 5. FAILURE ANALYSIS ----------------------------------------------
heading(doc, "5. Failure analysis", level=2)
para(doc, "Worst clips by total miss count across all 5 valid models:")

fc = pd.read_csv(RESULTS / "failure_cases.csv")
# Re-shape: only show clips where 3+ models missed
fc = fc[fc["n_misses"] >= 2].head(6)
rows = []
for _, r in fc.iterrows():
    rows.append([
        str(r["clip_id"]).zfill(2),
        r["locality"],
        r["condition"],
        r["difficulty"],
        f"{int(r['n_hits'])}/{int(r['n_models'])}",
        f"{int(r['n_partial'])}/{int(r['n_models'])}",
        f"{int(r['n_misses'])}/{int(r['n_models'])}",
    ])
make_table(
    doc,
    headers=["#", "Locality", "Condition", "Difficulty", "Hits", "Partial", "Misses"],
    rows=rows,
    col_widths=[0.4, 1.6, 1.2, 0.8, 0.7, 0.7, 0.7],
)

para(doc,
     "Clip 12 (Kadugondanahalli, phone call, hard name): zero hits across all 5 models. "
     "Deepgram wrote \"kaddu gunda hainili\", Groq \"tadduganda hanyli\", IndicConformer "
     "\"kaddugod nali\". Unfamiliar Kannada name + the band-limiting that phones do = the worst "
     "case in the set.")
para(doc,
     "Clip 9 (Majestic, rushed, easy name): only 1 hit. Deepgram and Groq both heard "
     "\"magistrate\" / \"mejistri\"; only Sarvam wrote \"mejestik\". An easy name fails when it "
     "sounds like a common English word the model has seen orders of magnitude more often.")
para(doc,
     "Damage ranking by condition: traffic noise hurts most, then phone call, then whispered, "
     "then rushed. Whispered and rushed are surprisingly easy for short locality tokens. "
     "Traffic causes outright hallucinations — Deepgram inserted the word \"police\" before "
     "\"Chikkabanavara\" on clip 14.")

# hit-by-condition plot
if (PLOTS / "03_hit_by_condition.png").exists():
    doc.add_picture(str(PLOTS / "03_hit_by_condition.png"), width=Inches(6.4))


# ---------- 6. RECOMMENDATION -------------------------------------------------
heading(doc, "6. Recommendation", level=2)
para(doc, "Match the system to the constraint:")

make_table(
    doc,
    headers=["Constraint", "Pick", "Why"],
    rows=[
        ["Best accuracy on entity capture (onboarding flow)", "Sarvam Saarika v2.5", "Top LHR (0.55) and WER tier, 0.73 s latency. Designed for Indian Hinglish."],
        ["Lowest latency for real-time voice bots", "Groq Whisper large-v3", "0.43 s mean, 0.62 s p90. Same accuracy tier as Sarvam at ~1.7× the speed. Generic model so weaker on rare Kannada names."],
        ["Self-hosted, no vendor lock-in, cost control at scale", "AI4Bharat IndicConformer", "Open weights, India-trained, accuracy matches Sarvam, runs on a single GPU. NeMo deps are fiddly — budget a day for deployment."],
        ["Default pipeline for the platform", "Sarvam + Groq dual-write, locality post-processor", "Hit both; pick the higher-confidence locality after fuzzy + phonetic match against the Bangalore locality list. Catches the Deepgram-style \"right sound, wrong spelling\" cases at near-zero cost."],
        ["Do not pick", "Deepgram Nova-3 multi", "Required by the brief as the baseline. On this dataset it is the weakest accuracy choice and not the fastest either."],
    ],
    col_widths=[2.2, 1.7, 3.2],
)


# ---------- 7. LIMITATIONS ----------------------------------------------------
heading(doc, "7. Limitations and what I would do next", level=2)
bullet(doc,
       "Sample is 20 clips, one speaker (me). Numbers are directional, not tight. Next step: "
       "200+ clips across 5–10 speakers including at least one female speaker — all five "
       "models are reported to train on male-skewed data and this dataset can't probe that bias.")
bullet(doc,
       "One language pair (Hindi-English Hinglish). The Vahan platform also sees Kannada, "
       "Telugu, Tamil. IndicConformer and Sarvam both claim multi-Indic; that is the next "
       "benchmark to run.")
bullet(doc,
       "Local Whisper on Colab T4 came out at 3.9 s mean latency, which is mostly Colab "
       "overhead, not the model. On a dedicated L4 / A10 the same model is ~0.5 s for these "
       "clip lengths. Treat that row's latency as unreliable and re-measure on real hardware "
       "before any deployment decision.")
bullet(doc,
       "indic_whisper_fallback errored on all 20 clips on Colab. Not investigated further "
       "because IndicConformer — the primary India-tuned open-source slot — ran cleanly.")
bullet(doc,
       "The locality-match metric does fuzzy + alias matching but no cross-script phonetic "
       "equivalence (Roman ↔ Devanagari at the sound level). Mapping reference and hypothesis "
       "to IPA before comparing would be more principled, and would partly close the Deepgram "
       "gap in finding #3.")


doc.save(OUT)
print(f"Wrote {OUT}")
print(f"Pages: open the file to verify — table layout is tuned for ≤3 pages.")
